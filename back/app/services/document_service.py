from io import BytesIO
from uuid import UUID

from docx import Document as DocxDocument

from app.dto.document import DocumentDownload
from app.entities.base import Document, DocumentContent
from app.repositories.document_content_repository import DocumentContentRepository
from app.repositories.document_repository import DocumentRepository
from app.services.rag_manager import RagManager


class DocumentService:
    types: dict[str, str] = {
        "txt": "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    def __init__(self, document_repository: DocumentRepository, content_repository: DocumentContentRepository, rag_manager: RagManager):
        self.document_repository = document_repository
        self.content_repository = content_repository
        self.rag_manager = rag_manager

    async def get_document_meta(self, document_id: UUID) -> Document:
        return await self.document_repository.get(document_id)

    async def get_documents_meta(self) -> list[Document]:
        return await self.document_repository.get_all()

    async def upload(self, filename: str, mime_type: str, content: bytes, encoding: str = "utf-8") -> Document:
        if mime_type == self.types["txt"]:
            text = content.decode(encoding)

        elif mime_type == self.types["docx"]:
            file = BytesIO(content)
            docx = DocxDocument(file)
            text = "\n".join(p.text for p in docx.paragraphs)

        else:
            raise ValueError("Unsupported file type")

        document = Document(
            filename=filename,
            mime_type=mime_type,
            size=len(content)
        )

        await self.document_repository.add(document)

        content = DocumentContent(
            document_id=document.id,
            encoding=encoding,
            text=text
        )

        await self.content_repository.add(content)

        self.rag_manager.start_rebuild()

        return document

    async def download(self, document_id: UUID) -> DocumentDownload:
        document = await self.document_repository.get(document_id)
        if not document:
            raise ValueError("Document not found")

        content = await self.content_repository.get_by_document(document_id)
        if not content:
            raise ValueError("Document content not found")

        if document.mime_type == self.types["txt"]:
            raw = content.text.encode(content.encoding)

        elif document.mime_type == self.types["docx"]:
            docx = DocxDocument()
            for line in content.text.split("\n"):
                docx.add_paragraph(line)

            buffer = BytesIO()
            docx.save(buffer)
            raw = buffer.getvalue()

        else:
            raise ValueError("Unsupported file type")

        return DocumentDownload(filename=str(document.filename), mime_type=str(document.mime_type), content=raw)

    async def delete(self, document_id: UUID) -> None:
        document = await self.document_repository.get(document_id)

        if not document:
            return

        content = await self.content_repository.get_by_document(document_id)

        if content:
            await self.content_repository.delete(content)

        await self.document_repository.delete(document)
        self.rag_manager.start_rebuild()
